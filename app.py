import streamlit as st
import os


# Set page config as the very first Streamlit command
st.set_page_config(
    page_title="AI Study Assistant",
    page_icon="🚀",
    layout="wide"
)

# ================================
# CONFIGURATION - API KEY HANDLING
# ================================
# Try to get API key from environment variable first
import os

try:
    GEMINI_API_KEY = st.secrets["GEMINI_API_KEY"]
except Exception:
    from dotenv import load_dotenv
    load_dotenv()
    GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

# ================================

import PyPDF2
import pdfplumber
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.colors import HexColor
import google.generativeai as genai
from datetime import datetime
import time
import hashlib

# Optimized Gemini initialization
@st.cache_resource
def init_ai():
    """Initialize AI with optimized settings"""
    api_key = GEMINI_API_KEY
    
    # If no environment variable, try getting from session state (manual entry)
    if not api_key and 'manual_api_key' in st.session_state:
        api_key = st.session_state.manual_api_key
    
    if not api_key:
        return None, "No API key provided"
    
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(
            'gemini-1.5-flash',
            generation_config={
                'temperature': 0.7,
                'top_p': 0.9,
                'top_k': 40,
                'max_output_tokens': 4096,
            }
        )
        return model, "ready"
    except Exception as e:
        return None, str(e)

class OptimizedStudyAssistant:
    def __init__(self):
        self.model, self.status = init_ai()
    
    @st.cache_data
    def extract_pdf_optimized(_self, file_bytes, file_hash):
        """Optimized PDF extraction with smart processing"""
        try:
            start_time = time.time()
            file_obj = io.BytesIO(file_bytes)
            
            # Try pdfplumber first for better quality
            try:
                with pdfplumber.open(file_obj) as pdf:
                    total_pages = len(pdf.pages)
                    text_chunks = []
                    
                    # Process in batches for better memory management
                    batch_size = 20
                    for i in range(0, min(total_pages, 200), batch_size):  # Limit to 200 pages max
                        batch_text = ""
                        end_page = min(i + batch_size, total_pages, 200)
                        
                        for page_num in range(i, end_page):
                            try:
                                page_text = pdf.pages[page_num].extract_text()
                                if page_text and page_text.strip():
                                    batch_text += page_text + "\n\n"
                            except:
                                continue
                        
                        if batch_text.strip():
                            text_chunks.append(batch_text)
                    
                    full_text = "\n".join(text_chunks)
                    pages_processed = min(total_pages, 200)
            
            except:
                # Fallback to PyPDF2
                file_obj.seek(0)
                pdf_reader = PyPDF2.PdfReader(file_obj)
                total_pages = len(pdf_reader.pages)
                text_chunks = []
                
                for i in range(min(total_pages, 200)):
                    try:
                        page_text = pdf_reader.pages[i].extract_text()
                        if page_text:
                            text_chunks.append(page_text)
                    except:
                        continue
                
                full_text = "\n\n".join(text_chunks)
                pages_processed = min(total_pages, 200)
            
            # Clean and optimize text
            full_text = re.sub(r'\n\s*\n', '\n\n', full_text)  # Remove excessive newlines
            full_text = re.sub(r'[^\w\s\.\,\!\?\:\;\-\(\)]', ' ', full_text)  # Remove special chars
            full_text = re.sub(r'\s+', ' ', full_text).strip()
            
            processing_time = time.time() - start_time
            
            return {
                'text': full_text,
                'pages_processed': pages_processed,
                'total_pages': total_pages,
                'char_count': len(full_text),
                'word_count': len(full_text.split()),
                'processing_time': round(processing_time, 2),
                'success': True
            }
            
        except Exception as e:
            return {
                'text': "",
                'error': str(e),
                'success': False
            }
    
    @st.cache_data
    def extract_docx_optimized(_self, file_bytes, file_hash):
        """Optimized DOCX extraction"""
        try:
            start_time = time.time()
            file_obj = io.BytesIO(file_bytes)
            doc = Document(file_obj)
            
            text_chunks = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_chunks.append(paragraph.text.strip())
            
            full_text = "\n\n".join(text_chunks)
            full_text = re.sub(r'\s+', ' ', full_text).strip()
            
            processing_time = time.time() - start_time
            
            return {
                'text': full_text,
                'char_count': len(full_text),
                'word_count': len(full_text.split()),
                'processing_time': round(processing_time, 2),
                'success': True
            }
        except Exception as e:
            return {
                'text': "",
                'error': str(e),
                'success': False
            }
    
    def smart_content_search(self, text, topic, max_content=8000):
        """Intelligent content search with relevance scoring"""
        if not topic or not text:
            return text[:max_content]
        
        # Tokenize topic
        topic_words = set(word.lower().strip() for word in re.findall(r'\b\w+\b', topic))
        topic_words = {word for word in topic_words if len(word) > 2}
        
        # Split into paragraphs and score them
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip() and len(p.strip()) > 50]
        
        scored_paragraphs = []
        for paragraph in paragraphs:
            para_words = set(word.lower() for word in re.findall(r'\b\w+\b', paragraph))
            
            # Calculate relevance score
            matches = len(topic_words.intersection(para_words))
            word_density = matches / max(len(para_words), 1)
            length_bonus = min(len(paragraph) / 500, 2)  # Prefer substantial paragraphs
            
            total_score = matches * 2 + word_density * 10 + length_bonus
            
            if total_score > 0:
                scored_paragraphs.append((total_score, paragraph))
        
        # Sort by relevance and combine top paragraphs
        scored_paragraphs.sort(reverse=True, key=lambda x: x[0])
        
        selected_content = []
        current_length = 0
        
        for score, paragraph in scored_paragraphs:
            if current_length + len(paragraph) <= max_content:
                selected_content.append(paragraph)
                current_length += len(paragraph)
            else:
                # Add partial paragraph if it fits
                remaining = max_content - current_length
                if remaining > 200:  # Only if substantial space left
                    selected_content.append(paragraph[:remaining] + "...")
                break
        
        result = "\n\n".join(selected_content)
        return result if result else text[:max_content]
    
    def generate_comprehensive_notes(self, text, topic):
        """Generate comprehensive notes from document content"""
        if not self.model:
            return f"# Error\n\n{self.status}"
        
        try:
            # Create a comprehensive prompt
            prompt = f"""You are an expert academic note-taker. Create comprehensive, well-structured study notes on "{topic}" based on the provided text.

Structure your notes as follows:

# {topic} - Study Notes

## Table of Contents
1. Overview & Introduction
2. Key Concepts & Definitions  
3. Main Topics & Details
4. Important Examples & Applications
5. Critical Points & Takeaways
6. Summary & Conclusions

## 1. Overview & Introduction
[Provide a clear introduction to the topic - 2-3 sentences]

## 2. Key Concepts & Definitions
[List and define the most important terms and concepts]
**Term 1**: Definition
**Term 2**: Definition

## 3. Main Topics & Details  
[Organize the main content into logical sections with subheadings]

### 3.1 [Subtopic 1]
• Key point 1
• Key point 2
• Key point 3

### 3.2 [Subtopic 2]
• Key point 1
• Key point 2

## 4. Important Examples & Applications
[Include relevant examples, use cases, or applications]

## 5. Critical Points & Takeaways
• Most important point 1
• Most important point 2
• Most important point 3
• Most important point 4
• Most important point 5

## 6. Summary & Conclusions
[Wrap up with a concise summary of the main ideas]

---

**Instructions:**
- Use clear, academic language suitable for studying
- Include specific details from the text
- Make definitions concise but complete
- Use bullet points for easy scanning
- Ensure logical flow between sections
- Focus specifically on "{topic}" while incorporating supporting information

**Source Text:**
{text}

Generate the comprehensive study notes following this exact structure:"""
            
            start_time = time.time()
            response = self.model.generate_content(prompt)
            generation_time = time.time() - start_time
            
            result = response.text
            result += f"\n\n---\n*Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')} in {round(generation_time, 1)}s*"
            
            return result
            
        except Exception as e:
            return f"# Error Generating Notes\n\n**Error Details:** {str(e)}\n\nPlease check your API key and internet connection."
    
    def generate_ai_knowledge_notes(self, topic, options):
        """Generate customized notes from AI's knowledge with extensive options"""
        if not self.model:
            return f"# Error\n\n{self.status}"
        
        try:
            # Extract options with defaults
            difficulty = options.get('difficulty', 'intermediate')
            note_style = options.get('note_style', 'comprehensive')
            focus_areas = options.get('focus_areas', [])
            length = options.get('length', 'detailed')
            format_style = options.get('format_style', 'academic')
            include_examples = options.get('include_examples', True)
            include_diagrams = options.get('include_diagrams', False)
            target_audience = options.get('target_audience', 'student')
            
            # Build dynamic prompt based on options
            if note_style == 'outline':
                structure = f"""
# {topic} - Study Outline

## I. Main Topic Overview
## II. Key Concepts (A, B, C format)
## III. Important Details
## IV. Summary Points"""
                
            elif note_style == 'flashcard':
                structure = f"""
# {topic} - Flashcard Format

## Front/Back Study Cards
**Card 1**: Question → Answer
**Card 2**: Term → Definition
**Card 3**: Concept → Explanation"""
                
            elif note_style == 'mind_map':
                structure = f"""
# {topic} - Mind Map Style

## Central Concept: {topic}
### Branch 1: [Main Category]
    • Sub-point 1
    • Sub-point 2
### Branch 2: [Main Category]
    • Sub-point 1
    • Sub-point 2"""
                
            elif note_style == 'qa':
                structure = f"""
# {topic} - Q&A Format

## Essential Questions & Detailed Answers
**Q1**: [Important question]?
**A1**: [Comprehensive answer]

**Q2**: [Important question]?
**A2**: [Comprehensive answer]"""
                
            else:  # comprehensive
                structure = f"""
# {topic} - Comprehensive Study Notes

## Table of Contents
1. Introduction & Overview
2. Key Concepts & Definitions
3. Core Topics & Detailed Analysis
4. Practical Applications & Examples
5. Important Facts & Figures
6. Common Misconceptions
7. Study Tips & Memory Aids
8. Advanced Topics & Further Learning"""
            
            # Build focus areas section
            focus_section = ""
            if focus_areas:
                focus_section = f"\n**Special Focus Areas**: {', '.join(focus_areas)}"
            
            # Length specifications
            length_guide = {
                'brief': "Keep each section concise (2-3 sentences or bullet points)",
                'standard': "Provide moderate detail (1-2 paragraphs per section)", 
                'detailed': "Include comprehensive explanations (2-4 paragraphs per section)",
                'extensive': "Provide thorough coverage with examples and elaboration"
            }
            
            # Format specifications
            format_guide = {
                'academic': "Use formal academic language and scholarly tone",
                'conversational': "Use friendly, easy-to-understand language",
                'technical': "Include technical terminology and precise definitions",
                'simplified': "Explain complex concepts in simple terms"
            }
            
            # Audience specifications
            audience_guide = {
                'student': "Structure for studying and exam preparation",
                'professional': "Focus on practical applications and industry relevance",
                'researcher': "Include latest developments and research directions",
                'general': "Make accessible to general audience with basic background"
            }
            
            # Examples and diagrams
            examples_text = ""
            if include_examples:
                examples_text = "Include specific, real-world examples and case studies."
            
            diagrams_text = ""
            if include_diagrams:
                diagrams_text = "Suggest where diagrams, charts, or visual aids would be helpful (describe what should be visualized)."
            
            prompt = f"""You are an expert educational content creator. Generate {format_style} study notes on "{topic}" for a {target_audience} audience.

**Content Specifications:**
- Difficulty Level: {difficulty}
- Note Style: {note_style}
- Length: {length} ({length_guide[length]})
- Format: {format_style} ({format_guide[format_style]})
- Target Audience: {target_audience} ({audience_guide[target_audience]}){focus_section}

{examples_text} {diagrams_text}

**Structure to Follow:**
{structure}

**Additional Instructions:**
- Write at {difficulty} level appropriate for {target_audience}
- Use {format_style} language style
- {length_guide[length]}
- Include practical study advice
- Make content engaging and memorable
- Focus on accuracy and educational value
- Structure information logically
- Use clear headings and bullet points
- Include key terms in **bold**

Generate comprehensive, well-structured study notes on "{topic}" following these specifications exactly."""
            
            start_time = time.time()
            response = self.model.generate_content(prompt)
            generation_time = time.time() - start_time
            
            result = response.text
            result += f"\n\n---\n*Generated with {note_style} style at {difficulty} level on {datetime.now().strftime('%B %d, %Y at %I:%M %p')} in {round(generation_time, 1)}s*"
            
            return result
            
        except Exception as e:
            return f"# Error Generating AI Notes\n\n**Error Details:** {str(e)}\n\nPlease check your API key and internet connection."
    
    def create_professional_docx(self, notes):
        """Create professional DOCX with proper formatting"""
        doc = Document()
        
        # Add custom styles safely
        styles = doc.styles
        
        # Create custom styles safely
        try:
            title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Calibri'
            title_style.font.size = Pt(18)
            title_style.font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass  # Style might already exist
        
        lines = notes.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if line.startswith('# '):
                doc.add_heading(line[2:], level=0)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=1)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=2)
            elif line.startswith('• ') or line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ')):
                doc.add_paragraph(line[3:], style='List Number')
            elif '**' in line:
                p = doc.add_paragraph()
                parts = line.split('**')
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1:
                        run.bold = True
            else:
                if line not in ['---', '___']:
                    doc.add_paragraph(line)
        
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io
    
    def create_professional_pdf(self, notes):
        """Create professional PDF with enhanced formatting"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        
        # Enhanced styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=20,
            alignment=TA_CENTER,
            spaceAfter=30,
            textColor=HexColor('#1f4e79')
        )
        
        heading1_style = ParagraphStyle(
            'CustomH1',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=12,
            spaceBefore=20,
            textColor=HexColor('#2e75b6')
        )
        
        heading2_style = ParagraphStyle(
            'CustomH2',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=8,
            spaceBefore=15,
            textColor=HexColor('#2e75b6')
        )
        
        story = []
        lines = notes.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 6))
                continue
            
            if line.startswith('# '):
                story.append(Paragraph(line[2:], title_style))
            elif line.startswith('## '):
                story.append(Paragraph(line[3:], heading1_style))
            elif line.startswith('### '):
                story.append(Paragraph(line[4:], heading2_style))
            elif line.startswith('• ') or line.startswith('- '):
                story.append(Paragraph(f"• {line[2:]}", styles['Normal']))
            elif '**' in line:
                # Properly handle bold formatting by replacing pairs of **
                formatted = line
                while '**' in formatted:
                    # Find first pair of ** and replace them
                    first_pos = formatted.find('**')
                    if first_pos != -1:
                        second_pos = formatted.find('**', first_pos + 2)
                        if second_pos != -1:
                            # Replace first ** with <b> and second ** with </b>
                            formatted = formatted[:first_pos] + '<b>' + formatted[first_pos+2:second_pos] + '</b>' + formatted[second_pos+2:]
                        else:
                            # If no matching **, just remove the **
                            formatted = formatted.replace('**', '')
                            break
                
                # Clean any remaining ** or malformed tags
                formatted = formatted.replace('**', '')
                story.append(Paragraph(formatted, styles['Normal']))
            else:
                if line not in ['---', '___'] and not line.startswith('*Generated'):
                    story.append(Paragraph(line, styles['Normal']))
        
        doc.build(story)
        buffer.seek(0)
        return buffer

def main():
    st.title("🚀 Advanced AI Study Assistant")
    st.markdown("**Secure & Customizable** • Environment-based API • Advanced note controls • Professional outputs")
    
    # API Key Configuration Section
    if not GEMINI_API_KEY:
        st.warning("⚠️ API Key Required")
        with st.expander("🔑 API Key Setup"):
            st.markdown("""
            **Two ways to set your API key:**
            
            **Option 1: Environment Variable (Recommended for security)**
            1. Create a `.env` file in your project folder
            2. Add: `GEMINI_API_KEY=your_actual_api_key_here`
            3. Restart the app
            
            **Option 2: Manual Entry (Temporary session)**
            """)
            
            manual_key = st.text_input(
                "Enter API Key (session only):",
                type="password",
                help="Get free API key from Google AI Studio"
            )
            
            if manual_key:
                st.session_state.manual_api_key = manual_key
                st.success("✅ API Key entered! You can now use the assistant.")
    
    # Initialize AI
    assistant = OptimizedStudyAssistant()
    
    if not assistant.model:
        st.error(f"⚠️ AI Configuration Error: {assistant.status}")
        if not GEMINI_API_KEY:
            st.info("📝 Get your free API key from [Google AI Studio](https://makersuite.google.com/app/apikey)")
        st.stop()
    
    st.success("✅ AI Ready - Secure Mode")
    
    # Session state initialization
    if 'notes' not in st.session_state:
        st.session_state.notes = ""
    if 'doc_info' not in st.session_state:
        st.session_state.doc_info = None
    if 'processing' not in st.session_state:
        st.session_state.processing = False
    if 'mode' not in st.session_state:
        st.session_state.mode = "document"
    if 'note_options' not in st.session_state:
        st.session_state.note_options = {}
    
    # Add mode selection at the top
    st.header("📚 Choose Your Study Mode")
    mode_col1, mode_col2 = st.columns(2)
    
    with mode_col1:
        if st.button("📄 Document-Based Notes", 
                    type="primary" if st.session_state.mode == "document" else "secondary",
                    use_container_width=True):
            st.session_state.mode = "document"
    
    with mode_col2:
        if st.button("🤖 Direct AI Knowledge", 
                    type="primary" if st.session_state.mode == "ai_direct" else "secondary",
                    use_container_width=True):
            st.session_state.mode = "ai_direct"
    
    # Show current mode
    if st.session_state.mode == "document":
        st.info("📄 **Document Mode**: Upload a document and generate notes from its content")
    else:
        st.info("🤖 **AI Knowledge Mode**: Ask AI directly about any topic using its training knowledge")
    
    st.markdown("---")
    col1, col2 = st.columns([1, 1])
    
    with col1:
        if st.session_state.mode == "document":
            # Document mode interface
            st.header("📚 Document Upload")
            
            # File uploader
            uploaded_file = st.file_uploader(
                "Upload your document",
                type=['pdf', 'docx', 'txt'],
                help="Supports large files • PDF up to 200 pages • Advanced processing"
            )
            
            # Process file automatically when uploaded
            if uploaded_file:
                # Create file hash for caching
                file_bytes = uploaded_file.read()
                file_hash = hashlib.md5(file_bytes).hexdigest()
                
                if st.session_state.doc_info is None or st.session_state.doc_info.get('hash') != file_hash:
                    with st.spinner("🔄 Processing document..."):
                        file_type = uploaded_file.name.split('.')[-1].lower()
                        
                        if file_type == 'pdf':
                            result = assistant.extract_pdf_optimized(file_bytes, file_hash)
                        elif file_type == 'docx':
                            result = assistant.extract_docx_optimized(file_bytes, file_hash)
                        elif file_type == 'txt':
                            text = file_bytes.decode('utf-8')
                            result = {
                                'text': text,
                                'char_count': len(text),
                                'word_count': len(text.split()),
                                'processing_time': 0.1,
                                'success': True
                            }
                        
                        if result['success']:
                            result['hash'] = file_hash
                            result['filename'] = uploaded_file.name
                            st.session_state.doc_info = result
                            
                            # Show processing stats
                            stats_col1, stats_col2 = st.columns(2)
                            with stats_col1:
                                st.metric("Characters", f"{result['char_count']:,}")
                                st.metric("Words", f"{result['word_count']:,}")
                            with stats_col2:
                                if 'pages_processed' in result:
                                    st.metric("Pages", f"{result['pages_processed']}")
                                st.metric("Process Time", f"{result['processing_time']}s")
                            
                            st.success("✅ Document processed successfully!")
                        else:
                            st.error(f"❌ Processing failed: {result.get('error', 'Unknown error')}")
                else:
                    # Show cached info
                    info = st.session_state.doc_info
                    st.info(f"📄 **{info['filename']}** - {info['word_count']:,} words - Cached")
            
            # Topic input for document mode
            st.header("🎯 Study Focus")
            topic = st.text_input(
                "What do you want to study from the document?",
                placeholder="e.g., 'Neural Networks', 'Photosynthesis', 'World War 2'",
                help="Be specific for better results from your document",
                key="doc_topic"
            )
            
            # Generate button for document mode
            if st.button("🤖 Generate Notes from Document", type="primary", use_container_width=True):
                if not st.session_state.doc_info:
                    st.error("Please upload a document first")
                elif not topic:
                    st.error("Please enter a study topic")
                else:
                    with st.spinner("🧠 AI analyzing document and generating notes..."):
                        try:
                            # Smart content search
                            relevant_content = assistant.smart_content_search(
                                st.session_state.doc_info['text'], topic
                            )
                            
                            # Generate comprehensive notes
                            notes = assistant.generate_comprehensive_notes(relevant_content, topic)
                            st.session_state.notes = notes
                            
                            st.success("✅ Document-based notes generated!")
                            
                        except Exception as e:
                            st.error(f"Generation failed: {str(e)}")
        
        else:
            # AI Direct mode interface with extensive customization
            st.header("🤖 AI Knowledge Studio")
            
            st.markdown("Create highly customized study notes from AI's vast knowledge base")
            
            # Main topic input
            topic = st.text_input(
                "📚 Study Topic:",
                placeholder="e.g., 'Machine Learning', 'Quantum Physics', 'French Revolution'",
                help="Enter any topic you want comprehensive notes about",
                key="ai_topic"
            )
            
            # Advanced customization panel
            st.subheader("🎛️ Note Customization")
            
            # Create tabs for different option categories
            tab1, tab2, tab3, tab4 = st.tabs(["📝 Style", "🎯 Content", "👥 Audience", "🔧 Advanced"])
            
            with tab1:
                st.markdown("**Note Style & Format**")
                col1, col2 = st.columns(2)
                
                with col1:
                    note_style = st.selectbox(
                        "Note Format:",
                        ["comprehensive", "outline", "flashcard", "qa", "mind_map"],
                        format_func=lambda x: {
                            "comprehensive": "📚 Comprehensive Guide",
                            "outline": "📋 Structured Outline", 
                            "flashcard": "🗃️ Flashcard Format",
                            "qa": "❓ Q&A Style",
                            "mind_map": "🧠 Mind Map Style"
                        }[x],
                        help="Choose how you want your notes structured"
                    )
                    
                    length = st.selectbox(
                        "Note Length:",
                        ["brief", "standard", "detailed", "extensive"],
                        index=2,
                        format_func=lambda x: {
                            "brief": "⚡ Brief Overview",
                            "standard": "📄 Standard Detail", 
                            "detailed": "📖 Detailed Coverage",
                            "extensive": "📚 Comprehensive Deep-Dive"
                        }[x]
                    )
                
                with col2:
                    format_style = st.selectbox(
                        "Writing Style:",
                        ["academic", "conversational", "technical", "simplified"],
                        index=0,
                        format_func=lambda x: {
                            "academic": "🎓 Academic Formal",
                            "conversational": "💬 Conversational Easy",
                            "technical": "⚙️ Technical Precise", 
                            "simplified": "🔍 Simplified Clear"
                        }[x]
                    )
                    
                    difficulty = st.selectbox(
                        "Difficulty Level:",
                        ["beginner", "intermediate", "advanced", "expert"],
                        index=1,
                        format_func=lambda x: {
                            "beginner": "🌱 Beginner Friendly",
                            "intermediate": "📈 Intermediate Level",
                            "advanced": "🔥 Advanced Depth",
                            "expert": "⭐ Expert Level"
                        }[x]
                    )
            
            with tab2:
                st.markdown("**Content Focus Areas**")
                
                focus_areas = st.multiselect(
                    "Special Focus (select multiple):",
                    [
                        "Historical Context", "Practical Applications", "Current Research",
                        "Key Definitions", "Mathematical Formulas", "Case Studies",
                        "Controversies & Debates", "Future Implications", "Related Fields",
                        "Common Mistakes", "Best Practices", "Industry Standards"
                    ],
                    help="Choose specific aspects to emphasize"
                )
                
                col1, col2 = st.columns(2)
                with col1:
                    include_examples = st.checkbox("Include Examples", value=True)
                    include_history = st.checkbox("Historical Context", value=True)
                    include_applications = st.checkbox("Real Applications", value=True)
                
                with col2:
                    include_diagrams = st.checkbox("Suggest Diagrams", value=False)
                    include_misconceptions = st.checkbox("Common Misconceptions", value=True)
                    include_study_tips = st.checkbox("Study Tips", value=True)
            
            with tab3:
                st.markdown("**Target Audience**")
                
                target_audience = st.selectbox(
                    "Primary Audience:",
                    ["student", "professional", "researcher", "general"],
                    format_func=lambda x: {
                        "student": "🎓 Student (Exam Prep)",
                        "professional": "💼 Professional (Work Application)",
                        "researcher": "🔬 Researcher (Academic Focus)",
                        "general": "👥 General Public (Basic Interest)"
                    }[x]
                )
                
                subject_area = st.selectbox(
                    "Subject Domain:",
                    ["General", "STEM", "Liberal Arts", "Business", "Medicine", "Engineering", 
                     "Social Sciences", "Technology", "Arts", "Law"],
                    help="Subject context for better terminology"
                )
            
            with tab4:
                st.markdown("**Advanced Options**")
                
                col1, col2 = st.columns(2)
                with col1:
                    creativity = st.slider("AI Creativity", 0.1, 1.0, 0.7, 0.1,
                                         help="Higher = more creative explanations")
                    
                    include_citations = st.checkbox("Suggest Citations", value=False)
                    include_timeline = st.checkbox("Timeline (if applicable)", value=False)
                
                with col2:
                    focus_practical = st.checkbox("Emphasize Practical Aspects", value=False)
                    include_prerequisites = st.checkbox("Prerequisites Section", value=True)
                    include_next_steps = st.checkbox("Next Learning Steps", value=True)
            
            # Compile options
            note_options = {
                'note_style': note_style,
                'difficulty': difficulty,
                'length': length,
                'format_style': format_style,
                'target_audience': target_audience,
                'focus_areas': focus_areas,
                'include_examples': include_examples,
                'include_history': include_history,
                'include_applications': include_applications,
                'include_diagrams': include_diagrams,
                'include_misconceptions': include_misconceptions,
                'include_study_tips': include_study_tips,
                'creativity': creativity,
                'subject_area': subject_area
            }
            
            # Generate button with preview of settings
            with st.expander("🔍 Preview Settings"):
                st.json(note_options)
            
            if st.button("🚀 Generate Custom AI Notes", type="primary", use_container_width=True):
                if not topic:
                    st.error("Please enter a topic to learn about")
                else:
                    with st.spinner(f"🤖 AI creating {note_style} style {difficulty} level notes on {topic}..."):
                        try:
                            notes = assistant.generate_ai_knowledge_notes(topic, note_options)
                            st.session_state.notes = notes
                            st.success("✅ Custom AI knowledge notes generated!")
                            
                        except Exception as e:
                            st.error(f"Generation failed: {str(e)}")
            
            # Quick presets
            st.subheader("⚡ Quick Presets")
            preset_cols = st.columns(4)
            
            presets = {
                "Exam Prep": {"note_style": "qa", "difficulty": "intermediate", "length": "detailed"},
                "Quick Review": {"note_style": "outline", "difficulty": "beginner", "length": "brief"},
                "Deep Study": {"note_style": "comprehensive", "difficulty": "advanced", "length": "extensive"},
                "Flashcards": {"note_style": "flashcard", "difficulty": "intermediate", "length": "standard"}
            }
            
            for i, (preset_name, preset_options) in enumerate(presets.items()):
                col = preset_cols[i % 4]
                with col:
                    if st.button(f"🎯 {preset_name}", key=f"preset_{i}", use_container_width=True):
                        if topic:
                            merged_options = {**note_options, **preset_options}
                            with st.spinner(f"Generating {preset_name} notes..."):
                                notes = assistant.generate_ai_knowledge_notes(topic, merged_options)
                                st.session_state.notes = notes
                                st.rerun()
                        else:
                            st.warning("Enter a topic first!")
            
            # Popular topics with one-click generation
            st.subheader("💡 Popular Topics")
            topic_cols = st.columns(3)
            
            popular_topics = [
                ("🤖 Artificial Intelligence", "artificial intelligence"),
                ("🌡️ Climate Change", "climate change"), 
                ("₿ Blockchain", "blockchain technology"),
                ("🧬 DNA & Genetics", "genetics and DNA"),
                ("🌍 World War II", "world war 2"),
                ("🎭 Shakespeare", "william shakespeare"),
                ("⚛️ Quantum Physics", "quantum physics"),
                ("🧠 Psychology", "human psychology"),
                ("💰 Economics", "basic economics")
            ]
            
            for i, (display_name, topic_name) in enumerate(popular_topics):
                col = topic_cols[i % 3]
                with col:
                    if st.button(display_name, key=f"topic_{i}", use_container_width=True):
                        with st.spinner(f"Generating notes on {topic_name}..."):
                            notes = assistant.generate_ai_knowledge_notes(topic_name, note_options)
                            st.session_state.notes = notes
                            st.rerun()

        # Advanced options (for document mode only)
        with st.expander("⚙️ Advanced Options"):
            if st.session_state.mode == "document":
                content_focus = st.selectbox(
                    "Content Focus:",
                    ["Comprehensive Overview", "Key Concepts Only", "Detailed Analysis", "Quick Summary"]
                )
                include_definitions = st.checkbox("Include Definitions", value=True)
                include_examples_adv = st.checkbox("Include Examples", value=True, key="advanced_include_examples")

    with col2:
        st.header("📝 Professional Study Notes")
        
        if st.session_state.notes:
            # Display notes
            st.markdown(st.session_state.notes)
            
            # Download section
            st.header("📥 Download Notes")
            
            # Create download files
            docx_file = assistant.create_professional_docx(st.session_state.notes)
            pdf_file = assistant.create_professional_pdf(st.session_state.notes)
            
            # Download buttons without nested columns
            st.download_button(
                "📄 Download DOCX",
                data=docx_file.getvalue(),
                file_name=f"study_notes_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            
            st.download_button(
                "📋 Download PDF",
                data=pdf_file.getvalue(),
                file_name=f"study_notes_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf",
                mime="application/pdf",
                use_container_width=True
            )
            
            # Clear notes button
            if st.button("🗑️ Clear Notes", use_container_width=True):
                st.session_state.notes = ""
                st.rerun()
            
            # Statistics
            word_count = len(st.session_state.notes.split())
            st.info(f"📊 Generated: {word_count:,} words of professional study notes")
            
        else:
            st.info("🎯 Choose a mode and generate notes to see them here")
            
            # Feature preview
            with st.expander("✨ What You'll Get"):
                st.markdown("""
                **Comprehensive Study Notes Include:**
                
                📋 **Table of Contents** - Clear structure
                📖 **Overview & Introduction** - Context setting
                🎯 **Key Concepts & Definitions** - Important terms
                📚 **Main Topics & Details** - Organized content
                💡 **Examples & Applications** - Real-world usage
                ✅ **Critical Takeaways** - Must-remember points
                📝 **Summary & Conclusions** - Wrap-up insights
                
                **Professional Features:**
                - Academic-quality writing
                - Logical information hierarchy
                - Easy-to-scan bullet points
                - Comprehensive coverage
                - Timestamped generation
                """)
    
    # Footer with performance info
    st.markdown("---")
    st.markdown("### 🚀 Performance Features:")
    perf_col1, perf_col2, perf_col3 = st.columns(3)
    
    with perf_col1:
        st.markdown("**🔄 Smart Caching**\nProcesses documents once, reuses results")
    
    with perf_col2:
        st.markdown("**🎯 Intelligent Search**\nFinds most relevant content automatically")
    
    with perf_col3:
        st.markdown("**📊 Large Document Support**\nHandles PDFs up to 200 pages efficiently")

if __name__ == "__main__":
    main()
